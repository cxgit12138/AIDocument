"""
Word文档转换代理
处理DOCX到其他格式的转换
"""

from docx import Document
from typing import Dict,Any


def convert_docx_to_txt(input_path: str, output_path: str):
    """将Word文档转换为纯文本"""
    doc = Document(input_path)
    text_content = '\n'.join([para.text for para in doc.paragraphs])

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text_content)


def convert_docx_to_html(input_path: str, output_path: str):
    """将Word文档转换为HTML，保留基本格式"""
    doc = Document(input_path)
    html_content = """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Converted Document</title>
    <style>
        body { font-family: Arial, sans-serif; }
        h1, h2, h3, h4, h5, h6 { margin-top: 1em; margin-bottom: 0.5em; }
        p { margin-top: 0.5em; margin-bottom: 0.5em; }
        ul, ol { margin-top: 0.5em; margin-bottom: 0.5em; }
        li { margin-bottom: 0.2em; }
    </style>
</head>
<body>
"""

    for para in doc.paragraphs:
        # 根据段落样式确定HTML标签
        style_name = para.style.name.lower() if para.style else ""

        if 'heading 1' in style_name or '标题 1' in style_name:
            html_content += f"<h1>{para.text}</h1>\n"
        elif 'heading 2' in style_name or '标题 2' in style_name:
            html_content += f"<h2>{para.text}</h2>\n"
        elif 'heading 3' in style_name or '标题 3' in style_name:
            html_content += f"<h3>{para.text}</h3>\n"
        elif 'heading 4' in style_name or '标题 4' in style_name:
            html_content += f"<h4>{para.text}</h4>\n"
        elif 'heading 5' in style_name or '标题 5' in style_name:
            html_content += f"<h5>{para.text}</h5>\n"
        elif 'heading 6' in style_name or '标题 6' in style_name:
            html_content += f"<h6>{para.text}</h6>\n"
        else:
            # 普通段落
            if para.text.strip():  # 非空段落
                html_content += f"<p>{para.text}</p>\n"
            else:  # 空段落
                html_content += "<br>\n"

    # 处理表格
    for table in doc.tables:
        html_content += "<table border='1' cellspacing='0' cellpadding='5'>\n"
        for row in table.rows:
            html_content += "<tr>\n"
            for cell in row.cells:
                html_content += f"<td>{cell.text}</td>\n"
            html_content += "</tr>\n"
        html_content += "</table>\n"

    html_content += "</body>\n</html>"

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_content)


def convert_docx_to_md(input_path: str, output_path: str):
    """将Word文档转换为Markdown，保留基本格式"""
    doc = Document(input_path)
    md_content = ""

    for para in doc.paragraphs:
        # 根据段落样式确定Markdown格式
        style_name = para.style.name.lower() if para.style else ""

        if 'heading 1' in style_name or '标题 1' in style_name:
            md_content += f"# {para.text}\n\n"
        elif 'heading 2' in style_name or '标题 2' in style_name:
            md_content += f"## {para.text}\n\n"
        elif 'heading 3' in style_name or '标题 3' in style_name:
            md_content += f"### {para.text}\n\n"
        elif 'heading 4' in style_name or '标题 4' in style_name:
            md_content += f"#### {para.text}\n\n"
        elif 'heading 5' in style_name or '标题 5' in style_name:
            md_content += f"##### {para.text}\n\n"
        elif 'heading 6' in style_name or '标题 6' in style_name:
            md_content += f"###### {para.text}\n\n"
        else:
            # 普通段落
            if para.text.strip():  # 非空段落
                # 简单处理粗体和斜体（需要更复杂的实现来处理run级别格式）
                md_content += f"{para.text}\n\n"
            else:  # 空段落
                md_content += "\n"

    # 处理表格
    for i, table in enumerate(doc.tables):
        if i > 0:
            md_content += "\n"
        # 创建表头
        if table.rows:
            header_row = table.rows[0]
            # 计算每列的最大宽度
            col_widths = []
            for cell in header_row.cells:
                col_widths.append(max(len(cell.text), 3))  # 最小宽度为3

            # 输出表头
            header_line = "|"
            separator_line = "|"
            for j, cell in enumerate(header_row.cells):
                width = col_widths[j]
                header_line += f" {cell.text.ljust(width)} |"
                separator_line += f" {'-' * width} |"

            md_content += header_line + "\n"
            md_content += separator_line + "\n"

            # 输出数据行
            for row in table.rows[1:]:
                data_line = "|"
                for j, cell in enumerate(row.cells):
                    width = col_widths[j]
                    data_line += f" {cell.text.ljust(width)} |"
                md_content += data_line + "\n"

        md_content += "\n"

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(md_content)


def convert_docx_to_pdf(input_path: str, output_path: str):
    """将Word文档转换为PDF（使用docx2pdf）"""
    try:
        from docx2pdf import convert
        convert(input_path, output_path)
    except Exception as e:
        raise Exception(f"转换为PDF失败: {str(e)}")


def extract_docx_metadata(input_path: str) -> Dict[str, Any]:
    """提取DOCX文档元数据"""
    doc = Document(input_path)

    # 获取核心属性
    core_properties = doc.core_properties

    metadata = {
        "title": core_properties.title or "",
        "author": core_properties.author or "",
        "created": core_properties.created.isoformat() if core_properties.created else "",
        "modified": core_properties.modified.isoformat() if core_properties.modified else "",
        "subject": core_properties.subject or "",
        "keywords": core_properties.keywords or "",
        "category": core_properties.category or "",
        "comments": core_properties.comments or "",
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "page_count": len(doc.sections)
    }

    return metadata


def convert_docx_with_formatting(input_path: str, output_path: str):
    """将Word文档转换为带有格式信息的HTML（保留更多格式）"""
    doc = Document(input_path)
    html_content = """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Converted Document</title>
    <style>
        body { font-family: Arial, sans-serif; }
        h1, h2, h3, h4, h5, h6 { margin-top: 1em; margin-bottom: 0.5em; }
        p { margin-top: 0.5em; margin-bottom: 0.5em; }
        .bold { font-weight: bold; }
        .italic { font-style: italic; }
        .underline { text-decoration: underline; }
        table { border-collapse: collapse; margin: 1em 0; }
        td, th { border: 1px solid #ccc; padding: 5px; }
    </style>
</head>
<body>
"""

    for para in doc.paragraphs:
        # 根据段落样式确定HTML标签
        style_name = para.style.name.lower() if para.style else ""

        # 确定段落标签
        if 'heading 1' in style_name or '标题 1' in style_name:
            tag = "h1"
        elif 'heading 2' in style_name or '标题 2' in style_name:
            tag = "h2"
        elif 'heading 3' in style_name or '标题 3' in style_name:
            tag = "h3"
        elif 'heading 4' in style_name or '标题 4' in style_name:
            tag = "h4"
        elif 'heading 5' in style_name or '标题 5' in style_name:
            tag = "h5"
        elif 'heading 6' in style_name or '标题 6' in style_name:
            tag = "h6"
        else:
            tag = "p"

        # 处理段落中的run（格式化文本）
        paragraph_html = f"<{tag}>"
        for run in para.runs:
            text = run.text
            if run.bold:
                text = f"<span class='bold'>{text}</span>"
            if run.italic:
                text = f"<span class='italic'>{text}</span>"
            if run.underline:
                text = f"<span class='underline'>{text}</span>"
            paragraph_html += text

        paragraph_html += f"</{tag}>\n"
        html_content += paragraph_html

    # 处理表格
    for table in doc.tables:
        html_content += "<table>\n"
        for i, row in enumerate(table.rows):
            tag = "th" if i == 0 else "td"
            html_content += "<tr>\n"
            for cell in row.cells:
                html_content += f"<{tag}>{cell.text}</{tag}>\n"
            html_content += "</tr>\n"
        html_content += "</table>\n"

    html_content += "</body>\n</html>"

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_content)