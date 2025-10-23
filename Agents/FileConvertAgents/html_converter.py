"""
HTML文件转换代理
处理HTML到其他格式的转换
"""

from bs4 import BeautifulSoup
import re
from typing import Dict, List


def convert_html_to_txt(input_path: str, output_path: str):
    """将HTML转换为纯文本"""
    with open(input_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'html.parser')

    # 移除script和style标签内容
    for script in soup(["script", "style"]):
        script.decompose()

    # 获取文本内容并清理空白字符
    txt_content = soup.get_text()

    # 清理多余的空白行
    lines = [line.strip() for line in txt_content.splitlines()]
    txt_content = '\n'.join(line for line in lines if line)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(txt_content)


def convert_html_to_md(input_path: str, output_path: str):
    """将HTML转换为Markdown，保留更多格式"""
    with open(input_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'html.parser')

    # 移除script和style标签内容
    for script in soup(["script", "style"]):
        script.decompose()

    md_content = ""

    # 处理标题
    for i in range(1, 7):
        for tag in soup.find_all(f'h{i}'):
            text = tag.get_text().strip()
            if text:
                md_content += f"{'#' * i} {text}\n\n"

    # 处理段落
    for p in soup.find_all('p'):
        text = p.get_text().strip()
        if text:
            # 处理段落内的内联格式
            formatted_text = _process_inline_formatting(p)
            md_content += f"{formatted_text}\n\n"

    # 处理列表
    for ul in soup.find_all('ul'):
        md_content += _convert_ul_to_md(ul)
        md_content += "\n"

    for ol in soup.find_all('ol'):
        md_content += _convert_ol_to_md(ol)
        md_content += "\n"

    # 处理链接
    for a in soup.find_all('a', href=True):
        text = a.get_text().strip()
        href = a['href']
        if text and href:
            md_content += f"[{text}]({href})\n\n"

    # 处理代码块
    for pre in soup.find_all('pre'):
        code = pre.get_text().strip()
        if code:
            md_content += f"\n{code}\n\n\n"

    # 处理行内代码
    for code in soup.find_all('code'):
        if code.parent.name != 'pre':  # 避免重复处理pre中的code
            text = code.get_text().strip()
            if text:
                md_content = md_content.replace(text, f"`{text}`")

    # 处理粗体和斜体
    for b in soup.find_all(['b', 'strong']):
        text = b.get_text().strip()
        if text:
            md_content = md_content.replace(text, f"**{text}**")

    for i in soup.find_all(['i', 'em']):
        text = i.get_text().strip()
        if text:
            md_content = md_content.replace(text, f"*{text}*")

    # 处理表格
    for table in soup.find_all('table'):
        md_content += _convert_table_to_md(table)
        md_content += "\n"

    # 处理图片
    for img in soup.find_all('img', src=True):
        alt = img.get('alt', '')
        src = img['src']
        md_content += f"![{alt}]({src})\n\n"

    # 清理多余的空行
    md_content = re.sub(r'\n{3,}', '\n\n', md_content)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

def _process_inline_formatting(element) -> str:
    """处理段落内的内联格式"""
    text = element.get_text()
    # 这里可以添加更复杂的内联格式处理逻辑
    return text.strip()

def _convert_ul_to_md(ul) -> str:
    """将无序列表转换为Markdown"""
    md_content = ""
    for li in ul.find_all('li'):
        text = li.get_text().strip()
        if text:
            md_content += f"- {text}\n"
    return md_content

def _convert_ol_to_md(ol) -> str:
    """将有序列表转换为Markdown"""
    md_content = ""
    for i, li in enumerate(ol.find_all('li'), 1):
        text = li.get_text().strip()
        if text:
            md_content += f"{i}. {text}\n"
    return md_content

def _convert_table_to_md(table) -> str:
    """将HTML表格转换为Markdown表格"""
    md_content = ""

    # 处理表头
    thead = table.find('thead')
    if thead:
        headers = thead.find_all('th')
        if headers:
            header_row = "|"
            separator_row = "|"
            for th in headers:
                text = th.get_text().strip()
                header_row += f" {text} |"
                separator_row += " --- |"
            md_content += f"{header_row}\n{separator_row}\n"

    # 处理表体
    tbody = table.find('tbody')
    rows = tbody.find_all('tr') if tbody else table.find_all('tr')

    for tr in rows:
        # 跳过表头行（如果在tbody中）
        if tr.find('th'):
            continue

        cells = tr.find_all(['td', 'th'])
        if cells:
            row_content = "|"
            for cell in cells:
                text = cell.get_text().strip()
                row_content += f" {text} |"
            md_content += f"{row_content}\n"

    return md_content


def convert_html_to_docx(input_path: str, output_path: str):
    """将HTML转换为DOCX文档，保持正确的元素顺序"""
    try:
        from docx import Document
        from docx.shared import Inches

        with open(input_path, 'r', encoding='utf-8') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        doc = Document()

        # 按照HTML中元素出现的顺序处理内容
        # 查找所有需要处理的标签，并按照它们在HTML中的顺序排序
        elements = []

        # 收集所有需要处理的元素
        for i in range(1, 7):
            for tag in soup.find_all(f'h{i}'):
                elements.append((tag.sourceline, 'heading', i, tag))

        for tag in soup.find_all('p'):
            elements.append((tag.sourceline, 'paragraph', None, tag))

        for tag in soup.find_all('ul'):
            elements.append((tag.sourceline, 'ulist', None, tag))

        for tag in soup.find_all('ol'):
            elements.append((tag.sourceline, 'olist', None, tag))

        # 按照在HTML中的顺序排序
        elements.sort(key=lambda x: x[0] if x[0] is not None else 0)

        # 按顺序处理元素
        for _, element_type, level, tag in elements:
            if element_type == 'heading':
                text = tag.get_text().strip()
                if text:
                    doc.add_heading(text, level=level)
            elif element_type == 'paragraph':
                text = tag.get_text().strip()
                if text:
                    doc.add_paragraph(text)
            elif element_type == 'ulist':
                for li in tag.find_all('li'):
                    text = li.get_text().strip()
                    if text:
                        doc.add_paragraph(text, style='List Bullet')
            elif element_type == 'olist':
                for li in tag.find_all('li'):
                    text = li.get_text().strip()
                    if text:
                        doc.add_paragraph(text, style='List Number')

        doc.save(output_path)

    except ImportError:
        raise Exception("需要安装python-docx库: pip install python-docx")

def extract_html_metadata(input_path: str) -> Dict[str, str]:
    """提取HTML文档元数据"""
    with open(input_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'html.parser')

    metadata = {
        "title": "",
        "description": "",
        "keywords": "",
        "author": "",
        "language": ""
    }

    # 提取标题
    title_tag = soup.find('title')
    if title_tag:
        metadata["title"] = title_tag.get_text().strip()

    # 提取meta标签信息
    for meta in soup.find_all('meta'):
        name = meta.get('name', '').lower()
        content = meta.get('content', '')

        if name == 'description':
            metadata["description"] = content
        elif name == 'keywords':
            metadata["keywords"] = content
        elif name == 'author':
            metadata["author"] = content
        elif meta.get('http-equiv', '').lower() == 'content-language':
            metadata["language"] = content

    return metadata