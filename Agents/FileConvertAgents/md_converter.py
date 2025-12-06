"""
Markdown文件转换代理
处理Markdown到其他格式的转换
"""

import markdown
from bs4 import BeautifulSoup
from typing import Dict, Any


def convert_md_to_html(input_path: str, output_path: str):
    """将Markdown转换为HTML，保留更多格式和样式"""
    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 使用扩展来支持更多Markdown语法
    md = markdown.Markdown(extensions=[
        'extra',  # 包含表格、围栏代码块等
        'codehilite',  # 代码高亮
        'toc',  # 目录
        'nl2br',  # 换行转<br>
    ])

    html_content = md.convert(md_content)

    # 添加基本的CSS样式
    full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Converted Document</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Helvetica, Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 24px;
            margin-bottom: 16px;
            font-weight: 600;
            line-height: 1.25;
        }}
        h1 {{ font-size: 2em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
        h2 {{ font-size: 1.5em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
        p {{ margin-top: 0; margin-bottom: 16px; }}
        a {{ color: #0366d6; text-decoration: none; }}
        a:hover {{ text-decoration: underline; }}
        code {{
            padding: 0.2em 0.4em;
            margin: 0;
            font-size: 85%;
            background-color: rgba(27,31,35,0.05);
            border-radius: 3px;
            font-family: 'SFMono-Regular', Consolas, 'Liberation Mono', Menlo, monospace;
        }}
        pre {{
            padding: 16px;
            overflow: auto;
            font-size: 85%;
            line-height: 1.45;
            background-color: #f6f8fa;
            border-radius: 3px;
        }}
        pre code {{
            padding: 0;
            margin: 0;
            font-size: 100%;
            background: transparent;
            border: 0;
        }}
        blockquote {{
            padding: 0 1em;
            color: #6a737d;
            border-left: 0.25em solid #dfe2e5;
            margin: 0 0 16px 0;
        }}
        table {{
            display: block;
            width: 100%;
            overflow: auto;
            border-collapse: collapse;
        }}
        table th, table td {{
            padding: 6px 13px;
            border: 1px solid #dfe2e5;
        }}
        table tr:nth-child(2n) {{
            background-color: #f6f8fa;
        }}
        ul, ol {{
            padding-left: 2em;
            margin-top: 0;
            margin-bottom: 16px;
        }}
        li {{
            margin-bottom: 0.25em;
        }}
        img {{
            max-width: 100%;
            box-sizing: content-box;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(full_html)


def convert_md_to_txt(input_path: str, output_path: str):
    """将Markdown转换为纯文本"""
    with open(input_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 使用markdown库转换为HTML，然后提取文本
    html = markdown.markdown(md_content, extensions=['extra', 'nl2br'])
    soup = BeautifulSoup(html, 'html.parser')
    txt_content = soup.get_text()

    # 清理多余的空白行
    lines = [line.strip() for line in txt_content.splitlines()]
    txt_content = '\n'.join(line for line in lines if line)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(txt_content)


def convert_md_to_docx(input_path: str, output_path: str):
    """将Markdown转换为DOCX文档"""
    try:
        from docx import Document
        from docx.shared import Inches
        import re

        with open(input_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        doc = Document()

        # 按行处理Markdown内容
        lines = md_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # 处理标题
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                if text:
                    doc.add_heading(text, level=level)
                i+=1
                continue

            # 处理无序列表
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                doc.add_paragraph(text, style='List Bullet')
                i+=1
                continue

            # 处理有序列表
            elif re.match(r'^\d+\.', line):
                text = re.sub(r'^\d+\.\s*', '', line).strip()
                doc.add_paragraph(text, style='List Number')
                i+=1
                continue

            # 处理代码块
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    doc.add_paragraph(code_text, style='Code')
                i+=1
                continue

            # 处理普通段落
            elif line:
                doc.add_paragraph(line)

            i += 1

        doc.save(output_path)

    except ImportError:
        raise Exception("需要安装python-docx库: pip install python-docx")
    except Exception as e:
        raise Exception(f"转换为DOCX失败: {str(e)}")


def extract_md_metadata(input_path: str) -> Dict[str, Any]:
    """提取Markdown文档元数据"""
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    metadata = {
        "title": "",
        "word_count": 0,
        "line_count": len(lines),
        "headers": [],
        "has_code_blocks": False,
        "has_tables": False,
        "has_images": False
    }

    # 提取标题（第一个#标题）
    for line in lines:
        if line.startswith('#'):
            metadata["title"] = line.lstrip('#').strip()
            break

    # 统计词数
    text_content = '\n'.join(lines)
    metadata["word_count"] = len(text_content.split())

    # 提取所有标题
    for line in lines:
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            metadata["headers"].append({"level": level, "text": text})

    # 检测特性
    metadata["has_code_blocks"] = '' in content
    metadata["has_tables"] = '|' in content and '\n|' in content
    metadata["has_images"] = '![' in content

    return metadata